from docx import Document
import base64
import uuid
import os
import requests
import re

def map_style_dependencies_with_text(document_path):
    """
    Analyzes style inheritance relationships and extracts text content,
    including text from embedded content (headers, footers, tables, text boxes).
    Requires: python-docx
    """
    doc = Document(document_path)
    text_content = []

    # Extract text content with styles from the main document
    for para in doc.paragraphs:
        para_text = para.text.strip()
        if para_text:
            text_content.append([
                para.style.name if para.style else 'Unknown',
                para_text,
                ''
            ])

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para_text = para.text.strip()
                    if para_text:
                        text_content.append([
                            para.style.name if para.style else 'Unknown',
                            para_text,
                            ''
                        ])

    # Extract text from headers and footers
    for section in doc.sections:
        for header in [section.header, section.footer]:
            if header:
                for para in header.paragraphs:
                    para_text = para.text.strip()
                    if para_text:
                        text_content.append([
                            para.style.name if para.style else 'Unknown',
                            para_text,
                            ''
                        ])

    return text_content


def replace_in_paragraphs(paragraphs, replacements):
    """
    Replaces text in paragraphs without removing images.
    """
    for paragraph in paragraphs:
        full_text = "".join(run.text for run in paragraph.runs).strip()
        for replacement in replacements:
            if not paragraph.style or paragraph.style.name != replacement['style']:
                continue

            target_text = replacement['text'].strip()
            translated_text = replacement['translated_text'].strip()

            # Normalize both target and translated text by removing non-alphanumeric characters
            target_alphanumeric = re.sub(r'[^a-zA-Z0-9]', '', target_text)
            translated_alphanumeric = re.sub(r'[^a-zA-Z0-9]', '', translated_text)

            # Update text without clearing the paragraph (preserving images)
            if target_alphanumeric == re.sub(r'[^a-zA-Z0-9]', '', full_text):
                # Modify text in runs instead of clearing
                remaining_text = translated_text
                for run in paragraph.runs:
                    xml = run._element
                    drawing = xml.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing")
                    if remaining_text:
                        run.text = remaining_text[:len(run.text)]  # Update part of text
                        remaining_text = remaining_text[len(run.text):]  # Remaining text
                    else:
                        run.text = ""  # Clear remaining runs
                    if drawing is not None:
                        print("not none")
                        run._element.append(drawing)
                        
            else:
                for run in paragraph.runs:
                    if run.text:
                        # Use regex to find and replace alphanumeric content
                        run.text = re.sub(target_alphanumeric, translated_alphanumeric, run.text)


def combined_replace(document_path, replacements):
    """
    Combines the structured replacement (for paragraphs, headers/footers, tables)
    with embedded content processing (tables and chart series) using the same
    list-of-dicts structure.
    
    Parameters:
      document_path: Path to the .docx file.
      replacements: A list of dictionaries where each dictionary has:
          'style': the style name (e.g., 'Normal', 'Heading 1')
          'text': the original text to be replaced
          'translated_text': the new text to replace with
          
    Returns:
      The path to the new document.
    """
    def convert_to_dict(nested_list):
        # Convert each inner list into a dictionary with the appropriate keys
        return [
            {
                'style': item[0],
                'text': item[1],
                'translated_text': item[2]
            }
            for item in nested_list
        ]
    
    replacements = convert_to_dict(replacements)
    doc = Document(document_path)
    
    # Replace in main document paragraphs.
    replace_in_paragraphs(doc.paragraphs, replacements)
    
    # Replace in tables (iterate over each cell's paragraphs).
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs, replacements)
    
    # Replace in headers and footers.
    for section in doc.sections:
        for part in (section.header, section.footer):
            if part:
                replace_in_paragraphs(part.paragraphs, replacements)
    
    # Process embedded content in inline shapes (charts).
    for shape in doc.inline_shapes:
        if hasattr(shape, 'chart'):
            chart_part = shape.chart
            # Translate each chart series' name.
            for series in chart_part.series:
                # Chart series don't have style info so we simply check for an exact text match.
                for replacement in replacements:
                    target_text = replacement['text'].strip()
                    translated_text = replacement['translated_text'].strip()
                    if hasattr(series, 'name') and series.name and series.name.strip() == target_text:
                        series.name = translated_text

    fn = str(uuid.uuid4())
    temp_save_path = f"/tmp/{fn}.docx"
    doc.save(temp_save_path)

    # Read and encode the document for transfer
    with open(temp_save_path, "rb") as f:
        doc_bytes = f.read()
        encoded_doc = base64.b64encode(doc_bytes).decode("utf-8")  # Convert to Base64 string

    # Python code to be executed remotely
    code = f"""
import uuid
import base64
from docx import Document

# Decode and save the file
doc_bytes = base64.b64decode("{encoded_doc}")
fn = uuid.uuid4()
fn = str(fn)
file_path = f"/tmp/sandbox/{{fn}}.docx"
with open(file_path, "wb") as f:
    f.write(doc_bytes)
    """

    SANDBOX_URL = os.getenv('SANDBOX_ENDPOINT', '')
    # Execute remotely and return result
    response = requests.post(f"{SANDBOX_URL}/upload", json={"code": code}, timeout=30)
    r = response.json().get("output", "No output received")
    if isinstance(r, dict):
        r = r["files"]
    return r
